#!/usr/bin/env python3
"""Generate Zeta Pharma Sales Enablement Guide Word document from markdown source."""

import re
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SOURCE = Path(__file__).resolve().parents[1] / "Zeta_Pharma_Sales_Enablement_Guide.md"
OUTPUT = Path(__file__).resolve().parents[1] / "Zeta_Pharma_Sales_Enablement_Guide.docx"


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Zeta Pharma Commercial Platform")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Complete Sales Enablement Guide")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline.add_run("Every Module · Every Capability · Business-Oriented")
    tag_run.font.size = Pt(12)
    tag_run.italic = True
    tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"June 2026\nFor Sales Professionals & Commercial Leaders")
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < cols:
                table.rows[i].cells[j].text = cell_text
    doc.add_paragraph()


def add_formatted_paragraph(doc: Document, text: str, style: Optional[str] = None) -> None:
    """Add paragraph with basic bold/italic from markdown markers."""
    # Strip markdown links [text](url) -> text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Strip inline code
    text = text.replace("`", "")

    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def convert_markdown_to_docx(md_path: Path, output_path: Path) -> None:
    doc = Document()
    set_document_defaults(doc)
    add_title_page(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code_block = False
    table_rows: list[list[str]] = []
    skip_until = 0

    while i < len(lines):
        line = lines[i]

        # Skip content before first real chapter (after title page we add from # Part 0)
        if i < 5 and line.startswith("# Zeta Pharma"):
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block and "mermaid" not in lines[i - 1]:
                pass
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # Table handling
        if line.strip().startswith("|"):
            if is_table_separator(line):
                i += 1
                continue
            table_rows.append(parse_table_row(line))
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            add_formatted_paragraph(doc, text, style="List Number")
        else:
            add_formatted_paragraph(doc, stripped)

        i += 1

    if table_rows:
        add_table(doc, table_rows)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"\n— End of Document —\nZeta Pharma · {date.today():%B %Y}"
    )
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    print(f"Created: {output_path}")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source not found: {SOURCE}")
    convert_markdown_to_docx(SOURCE, OUTPUT)


if __name__ == "__main__":
    main()
