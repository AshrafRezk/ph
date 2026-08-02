#!/usr/bin/env python3
"""Generate Zeta Pharma product overview Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "Zeta_Pharma_Product_Overview.docx"


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Zeta Pharma Commercial Platform")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Product Overview & Feature Guide")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"June 2026\nBusiness & Features Orientation")
    meta_run.font.size = Pt(12)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_intro(doc: Document) -> None:
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "The Zeta Pharma Commercial Platform brings together everything a pharmaceutical "
        "field force and its leadership need to plan, execute, measure, and improve "
        "customer engagement. Built on Salesforce, the solution is delivered through "
        "two complementary experiences:"
    )
    bullets = [
        (
            "Field App — designed for medical representatives and MSLs in the field. "
            "It puts daily priorities, visit planning, call reporting, CLM presentations, "
            "and account intelligence at their fingertips — on desktop, tablet, and mobile."
        ),
        (
            "Management App — designed for district managers, sales force effectiveness "
            "(SFE), and C-level leadership. It provides organization-wide KPIs, workforce "
            "visibility, team performance drill-down, and a central Admin Console for "
            "configuring territories, products, CLM content, coaching templates, and plan cycles."
        ),
    ]
    for text in bullets:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph(
        "This document walks through each major module from a business perspective — "
        "what problem it solves, what value it delivers, and which capabilities are "
        "available today. Screenshot placeholders are included throughout; replace each "
        "placeholder with your own captures before sharing externally."
    )


def add_screenshot_placeholder(doc: Document, label: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ INSERT SCREENSHOT: {label} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)

    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box_run = box.add_run(" ")
    box_run.font.size = Pt(4)
    doc.add_paragraph()


def add_section(doc: Document, heading: str, business: str, features: list[str], screenshot: str) -> None:
    doc.add_heading(heading, level=2)

    doc.add_heading("Business Value", level=3)
    doc.add_paragraph(business)

    doc.add_heading("Key Features", level=3)
    for feature in features:
        doc.add_paragraph(feature, style="List Bullet")

    add_screenshot_placeholder(doc, screenshot)


def build_document() -> Document:
    doc = Document()
    set_document_defaults(doc)
    add_title_page(doc)
    add_intro(doc)

    # ── Field App ──────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Part 1 — Field App", level=1)
    doc.add_paragraph(
        "The Field App is the daily workspace for medical representatives. "
        "It follows industry-standard patterns from OCE and Veeva — home dashboard, "
        "weekly planner, structured call reporting, CLM presentations, and rich "
        "HCP account profiles — optimized for speed in the field."
    )

    add_section(
        doc,
        "Homepage",
        "Field reps start every day on a single landing page that answers three questions: "
        "How am I performing against plan? What should I do today? Who should I call next? "
        "The homepage replaces scattered reports and spreadsheets with live, territory-scoped "
        "KPIs and actionable recommendations — reducing time spent on admin and increasing "
        "time with customers.",
        [
            "Performance rings — month-to-date visit coverage % and call frequency (RF%) overall and by A/B/C classification, with drill-down to individual accounts.",
            "Gamification — activity streaks, milestone badges, and district leaderboard rankings to drive healthy competition.",
            "Today's Plan — ordered route stops for the day with an interactive map, drive-time estimates, and one-click navigation to Google Maps.",
            "Route optimization ideas — suggestions to improve daily coverage based on geography and visit targets.",
            "Top 5 Next Best Customers (NBC) — ranked recommendations with call-plan context and a quick action to create a draft visit.",
            "Mobile-first layout — fully responsive on Salesforce mobile for reps working from phones and tablets.",
        ],
        "Field App — Homepage dashboard",
    )

    add_section(
        doc,
        "Planner",
        "Effective field teams plan their week deliberately. The Planner gives reps a "
        "Salesforce-calendar-style weekly view where they can schedule visits, block time "
        "off, and visualize their day's route on a map — all without leaving the app.",
        [
            "Week calendar (Mon–Sun) with 30-minute time slots from 6:00 AM to midnight.",
            "Drag-and-drop scheduling — pull an HCP account from the sidebar onto a time slot to create a planned visit instantly.",
            "Reschedule and resize — move visit blocks to new days/times or drag the bottom handle to change duration.",
            "Time Off Territory (TOT) — drag a TOT block onto the calendar to submit leave requests that appear alongside visits.",
            "Account sidebar with filters — search and filter territory accounts by record type, specialty, classification, and brick.",
            "Map & Route view — see geocoded visit pins for any day, build a driving route across ordered stops, and view distance and duration.",
            "Route optimization — OSRM-powered routing with alternative route options; apply optimized times back to the calendar.",
            "Manager view — district managers can switch context to view and coach on a rep's plan.",
        ],
        "Field App — Weekly Planner (calendar + map view)",
    )

    add_section(
        doc,
        "Visit / Call Report",
        "Every customer interaction must be captured accurately and compliantly. The Visit "
        "module is the structured call report — aligned with OCE/Veeva call reporting "
        "terminology — where reps log visit details, attendees, products discussed, samples "
        "distributed, and CLM presentations used.",
        [
            "Section-based call report — Details, Attendees, Products, Samples, Presentations, and Affiliations in a left-nav shell (desktop) or mobile Visualforce page (Salesforce One).",
            "Attendee management — add the primary HCP plus affiliated contacts discovered via account affiliations or territory search.",
            "Product discussion capture — select territory-aligned products, record topic messages with sentiment (positive / neutral / negative).",
            "Sample distribution — issue samples from rep inventory with lot tracking; inventory is validated and deducted on visit completion.",
            "Visit lifecycle — Draft → In Progress → Completed with read-only enforcement after submission.",
            "Double Visit support — flag and link coaching events when a manager accompanies the rep.",
            "Quick actions — jump to account profile, start coaching evaluation, raise a medical inquiry, or send WhatsApp messages directly from the visit header.",
        ],
        "Field App — Visit / Call Report",
    )

    add_section(
        doc,
        "WhatsApp Reminder & Feedback",
        "Reps engage HCPs where they already communicate. WhatsApp integration lets reps "
        "send meeting reminders before a visit and product feedback surveys after — "
        "capturing HCP sentiment without paper forms or follow-up calls.",
        [
            "Meeting reminder — compose a pre-filled WhatsApp message with visit date, time, and location; open directly in WhatsApp to the HCP's phone number.",
            "Product survey — send a personalized survey link covering territory-aligned products discussed during the visit.",
            "Phone number discovery — automatically pulls mobile numbers from the visit account and attendee records.",
            "Feedback capture — HCP responses flow back into Visit Product Survey Feedback records, tagged with source \"WhatsApp Survey\".",
            "Message preview — review the full message body before sending to ensure accuracy.",
        ],
        "Field App — WhatsApp reminder and product survey",
    )

    add_section(
        doc,
        "CLM (Closed Loop Marketing)",
        "Approved presentations are the backbone of compliant detailing. CLM lets reps "
        "launch territory-targeted slide decks during a visit, automatically logs time on "
        "each slide, and captures HCP reactions to product messages — giving marketing and "
        "medical affairs real-world evidence of message effectiveness.",
        [
            "Presentation library — reps see only presentations published and targeted to their territory.",
            "Fullscreen slide player — launch from the Visit Presentations tab; pause/resume tracking as the conversation flows.",
            "Per-slide dwell time — every second on each slide is logged to the visit session for compliance and analytics.",
            "Product-message feedback — capture Positive / Neutral / Negative reactions per topic during or after the presentation.",
            "Configurable rating layouts — admin-defined forms for structured HCP response capture.",
            "Activity history — account timeline shows which CLMs were used in past visits for continuity across calls.",
        ],
        "Field App — CLM presentation player during visit",
    )

    add_section(
        doc,
        "Accounts",
        "Reps need a 360° view of every customer — HCP, pharmacy, institution, or business "
        "contact — with the right fields, related activity, and segmentation data for their "
        "territory. The Accounts module provides OCE/Veeva-style account profiles tailored "
        "to pharmaceutical record types.",
        [
            "Four record types — Medical Professional (HCP), Pharmacy, Institution (HCO), and Business Contact (BC), each with dedicated layouts and Lightning record pages.",
            "HCP enrichment — up to four medical specialties, designation, license, tier, and classification fields.",
            "Tabbed account profile — Overview, Visits, Samples, Affiliations, and Activity timeline on a single record page.",
            "Territory-scoped visibility — reps see accounts aligned to their territory via Enterprise Territory Management.",
            "Visit history — chronological list of past and upcoming visits with status and products discussed.",
            "Sample transaction history — audit trail of all sample movements linked to the account.",
        ],
        "Field App — HCP Account detail page",
    )

    add_section(
        doc,
        "Sell-Out Data & Pharmacy Matching",
        "Pharmaceutical sell-out from distributors is the starting point for market intelligence. "
        "Wholesaler feeds from IbnSina and Pharmaoverseas are loaded into Salesforce through a "
        "structured mapping regimen that matches each withdrawal row to the correct pharmacy "
        "account — without duplicates — using external identifiers, brick geography, and "
        "row-level validation before any data is committed.",
        [
            "Distributor import — upload wholesaler CSV files via Admin Console → Sales Data; supports IbnSina and Pharmaoverseas as data sources.",
            "Pre-import validation — preview rows before commit; required fields, allowed data sources, valid report months, known pharmacy and product IDs, and positive quantities are enforced.",
            "Pharmacy matching — each row's pharmacy external ID resolves to a single Salesforce pharmacy account, ensuring sell-out is attributed to the right outlet with no duplicate accounts.",
            "Location alignment — IQVIA IMS bricks (governorate, city, territory) assign each pharmacy to the correct geographic market cell; brick membership is maintained without duplicate links.",
            "Dedup on load — unique keys on withdrawal rows and brick–pharmacy memberships prevent duplicate sell-out or membership records when files are re-imported.",
            "Brick auto-assignment — withdrawal records inherit the pharmacy's active brick membership, keeping analytics aligned to territory geography.",
            "Downstream action — validated sell-out data powers pharmacy sales dashboards, ROI views, and planning recommendations that connect rising pharmacy performance to affiliated HCPs.",
        ],
        "Sell-out data import and pharmacy matching (Admin Console — Sales Data)",
    )

    add_section(
        doc,
        "Affiliations",
        "Affiliations bridge distributor sell-out signals to HCP engagement. While pharmacy "
        "withdrawal data shows where product is moving, affiliations reveal who influences "
        "those outlets — connecting rising pharmacy performance to the HCPs reps should "
        "prioritize. Reps still get a full relationship network for visit planning; "
        "leadership gets a commercial link between market data and field action.",
        [
            "Sell-out-informed insights — top-performing pharmacies from withdrawal analytics are cross-referenced with affiliation records to surface affiliated HCPs for visit recommendations.",
            "Pharmacy–HCP relationships — Partner, Pharmacist, and other affiliation types link pharmacies to the professionals who influence prescribing and purchasing decisions.",
            "Planning recommendations — when pharmacy sell-out rises in a territory, the system can recommend scheduling visits to affiliated accounts (e.g. engage the HCP linked to a high-growth pharmacy).",
            "Interactive network graph — visual tree of affiliated accounts with record-type icons and relationship types on account record pages.",
            "Click-to-navigate — select any node in the network and jump directly to that account's profile.",
            "Account record page tab — dedicated Affiliations tab on HCP, Pharmacy, Institution, and BC record pages.",
            "Visit integration — view affiliated contacts and add them as attendees directly from the call report Affiliations section.",
            "Mobile list view — collapsible affiliation list on smaller screens when the graph is not practical.",
        ],
        "Field App — Account Affiliations network view",
    )

    add_section(
        doc,
        "Coaching Events",
        "Structured field coaching turns manager ride-alongs into measurable development. "
        "Coaching Events link a competency template to a specific rep–manager session, "
        "with dual scoring that reveals alignment gaps and development priorities.",
        [
            "Template-driven assessments — competency questions grouped by section (Core Values, Selling Skills, etc.) on a gradient scale.",
            "Dual scoring — both the rep and the manager score each competency; the system computes section totals, strengths, weaknesses, and calibration gaps.",
            "Double Visit linkage — coaching events can be created and evaluated directly from a visit marked as a Double Visit.",
            "Evaluation workflow — Draft → In Progress → Review → Completed, with manager review gate.",
            "Mobile-compatible evaluation UI — complete assessments on tablet during or immediately after the field ride-along.",
            "Score progression — historical coaching scores feed into management dashboards for trend analysis.",
        ],
        "Field App — Coaching Event evaluation (double visit)",
    )

    # ── Management App ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Part 2 — Management App", level=1)
    doc.add_paragraph(
        "The Management App gives district managers, SFE teams, and C-level executives "
        "the visibility and control they need to run the commercial operation. From a "
        "single executive home page through team KPI drill-down to the Admin Console, "
        "leaders can monitor performance, identify gaps, and configure the platform "
        "without IT involvement."
    )

    add_section(
        doc,
        "Homepage (Executive Overview)",
        "Leadership needs the big picture first. The Executive Overview homepage rolls up "
        "organization-wide metrics across business units — headcount, field activity, CLM "
        "adoption, and visit trends — with the ability to drill into any BU, line, or "
        "district.",
        [
            "Organization Snapshot — clickable KPI tiles for total headcount, active field reps, managers, monthly visits, and CLM session count.",
            "6-month visit trend chart — bar chart of completed field visits across all reps for quick trajectory assessment.",
            "BU performance cards — per-business-unit rollup with workforce count, visit rate, coverage %, and CLM adoption.",
            "Hierarchy drill-down — navigate from BU → Line → District to scope all metrics to the level that matters.",
            "Workforce panel — click headcount tiles to open the full roster of managers or field reps for the selected scope.",
            "Top Performers — highlight best-in-class reps and managers by coverage, CLM usage, or coaching score for the selected scope.",
        ],
        "Management App — Executive Overview homepage",
    )

    add_section(
        doc,
        "Workforce",
        "People are the largest investment in any field force. The Workforce module gives "
        "managers and executives a live roster with performance context — connecting "
        "headcount data to the KPIs that matter for each employee.",
        [
            "Roster by role — view all Field Reps or Managers filtered by BU, line, or district.",
            "Employee detail on click — select any row to see attendance, compensation, and individual KPI summary.",
            "Top Performers grid — automatically surfaces leaders in coverage, visit rate, CLM adoption, and coaching scores.",
            "Payroll visibility — total monthly base payroll rollup for the selected workforce scope (executive view).",
            "Team KPI Command Center — dedicated dashboard with workforce tiles (headcount, monthly visits, CLM %, not-reported days, TOT days) filterable by BU/line/district.",
            "Medical Rep 360 — rep-level dashboard with submitted calls, coverage by rating, frequency status (LCF/RCF/MCF), visit targets, and coaching score progression.",
        ],
        "Management App — Workforce roster and Team KPI dashboard",
    )

    add_section(
        doc,
        "Admin Console",
        "Commercial operations change every cycle — new products launch, territories shift, "
        "presentations update, and plan targets reset. The Admin Console is the self-service "
        "configuration hub where authorized admins manage all platform modules without "
        "deployments or developer support.",
        [
            "CLM Management — upload PDF/HTML/ZIP presentations, configure slide sequences, link slides to products and messages, set territory targeting, and publish.",
            "Rating Layouts — design account, territory, and product rating forms with live preview for HCP feedback capture.",
            "Coaching Management — browse, search, create, and edit coaching templates with section-based competency questions.",
            "Territory Management — manage product lines, edit territory hierarchy, assign users, and seed demo field-force accounts.",
            "Bricks Management — define IQVIA IMS bricks, align them to territories, and manage pharmacy account membership.",
            "Products Manager — browse the product catalog by brand and align products to territory hierarchies.",
            "Plan Manager — manage monthly plan cycles, review employee coverage targets, and copy plans between months.",
            "Sales Data — import wholesaler sell-out CSVs from IbnSina and Pharmaoverseas with validation preview; pharmacy external-ID matching, brick auto-assignment, and deduplicated withdrawal loading; review import batches and loaded data.",
            "Integrations Management — monitor connectors for IMS Health, OneKey, Maps, Mendix, and other external platforms.",
        ],
        "Management App — Admin Console module tiles",
    )

    # ── Closing ────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Together, the Field App and Management App form a closed commercial loop:"
    )
    loop = [
        "Leaders configure territories, products, CLM content, coaching templates, and plan cycles in the Admin Console.",
        "Reps plan their week, execute visits, present CLM slides, capture feedback, and log calls in the Field App.",
        "WhatsApp reminders and surveys extend engagement beyond the face-to-face visit.",
        "Distributor sell-out data is matched to the right pharmacies; affiliations connect rising pharmacy performance to the HCPs reps should engage.",
        "Coaching events develop rep competency with measurable dual-score assessments.",
        "Management dashboards roll up performance in real time — enabling data-driven decisions on budget, CLM focus, and coaching priorities.",
    ]
    for item in loop:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Replace the screenshot placeholders throughout this document with your own "
        "captures, then share with stakeholders for demos, onboarding, or commercial "
        "cycle planning sessions."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"\n— End of Document —\nZeta Pharma · {date.today():%B %Y}")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
