#!/usr/bin/env python3
"""Generate Zeta Pharma technical proposal Word document (4 pages)."""

from datetime import date
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "Zeta_Pharma_Technical_Proposal.docx"

NAVY = RGBColor(0x1A, 0x36, 0x5D)
BLUE = RGBColor(0x44, 0x72, 0xC4)
GRAY = RGBColor(0x66, 0x66, 0x66)

# Each module: name + list of (persona_label, enablement_text)
MODULES = [
    {
        "name": "Home Page",
        "enablements": [
            ("reps", "view territory-scoped KPIs (coverage %, RF% by A/B/C), today's route map, and Top 5 Next Best Customers"),
            ("managers", "access team KPI tiles and workforce headcount from Management Home with BU → Line → District drill-down"),
            ("executives", "see organization-wide snapshots, 6-month visit trends, and BU performance cards on Executive Home"),
        ],
    },
    {
        "name": "Planner",
        "enablements": [
            ("reps", "plan the week on a drag-and-drop calendar with account sidebar, map pins, and OSRM route optimization"),
            ("managers", "view and coach on a rep's weekly plan and scheduled visits"),
        ],
    },
    {
        "name": "Time Off Territory",
        "enablements": [
            ("reps", "submit leave requests from the planner palette, integrated with KPI and time-card calculations"),
            ("managers", "approve, reject, or bulk-process TOT requests with overlap validation"),
        ],
    },
    {
        "name": "Visit / Call Report",
        "enablements": [
            ("reps", "log structured HCP visits (attendees, products, messages, samples, CLM) on desktop and mobile"),
            ("managers", "review submitted calls and double-visit coaching linkage in MR 360 reports"),
            ("executives", "link field visits to strategic launch projects for closed-loop tracking"),
            ("MSLs", "trace products discussed, CLM sessions, and inquiries back to source visits"),
            ("product managers", "connect visit execution to launch projects and promo budget lines"),
        ],
    },
    {
        "name": "Sample Distribution",
        "enablements": [
            ("reps", "issue lot-tracked samples from inventory with validation; auto-deduct on visit completion"),
            ("managers", "audit sample transactions and compliance via account and rep reports"),
        ],
    },
    {
        "name": "CLM (Closed Loop Marketing)",
        "enablements": [
            ("reps", "launch territory-targeted presentations with per-slide dwell-time and HCP message feedback"),
            ("managers", "monitor CLM adoption % and presentation usage across the team"),
            ("executives", "review CLM session counts and message effectiveness rollups"),
            ("MSLs", "govern scientific content usage via CLM adoption and presentation analytics"),
            ("product managers", "upload, sequence, target, and publish approved presentations via Admin Console"),
        ],
    },
    {
        "name": "WhatsApp Reminder & Product Survey",
        "enablements": [
            ("reps", "send meeting reminders and post-visit product survey links to HCPs via WhatsApp"),
            ("product managers", "track HCP survey sentiment and message feedback by product and rep"),
        ],
    },
    {
        "name": "Accounts",
        "enablements": [
            ("reps", "view HCP, Pharmacy, Institution, and Business Contact profiles with ATF/ATPF segmentation and visit history"),
            ("managers", "review account coverage and classification distribution by territory"),
            ("MSLs", "access consolidated HCP and Institution profiles with specialty depth and activity timeline"),
        ],
    },
    {
        "name": "Accounts Management",
        "enablements": [
            ("reps", "search and filter territory accounts from the Accounts Tab by record type, specialty, classification, and brick"),
            ("managers", "manage territory account alignment, ATF/ATPF classifications, and field-force account seeding via Admin Console"),
            ("product managers", "design account and territory rating layouts that feed ATF/ATPF segmentation forms"),
        ],
    },
    {
        "name": "Affiliations",
        "enablements": [
            ("reps", "view pharmacy–HCP relationship networks on account record pages and add affiliated contacts as visit attendees"),
            ("managers", "connect rising pharmacy sell-out to affiliated HCPs for visit recommendations"),
            ("executives", "link market data to field action via affiliation-informed planning insights"),
        ],
    },
    {
        "name": "Affiliations Management",
        "enablements": [
            ("reps", "create and maintain Account Affiliation records (type, role, strength) linking HCPs, pharmacies, and institutions"),
            ("managers", "filter affiliation networks by role, type, strength, inactive status, and outside-territory flags"),
            ("MSLs", "manage HCP–institution and influencer relationships for scientific engagement mapping"),
        ],
    },
    {
        "name": "Coaching Events",
        "enablements": [
            ("reps", "participate in dual-score competency evaluations during double visits"),
            ("managers", "conduct, score, and review field coaching events with gap analysis and score trends"),
            ("product managers", "configure coaching templates and competency sections via Admin Console"),
        ],
    },
    {
        "name": "Medical Inquiry",
        "enablements": [
            ("reps", "raise HCP medical questions to Medical Affairs directly from a visit"),
            ("MSLs", "receive, respond to, and close inquiries routed via the Medical Affairs Case queue"),
        ],
    },
    {
        "name": "Management Home",
        "enablements": [
            ("managers", "monitor team KPIs, Egypt coverage heat maps, and HQ announcements with hierarchy drill-down"),
            ("executives", "access the same workforce and KPI command center scoped to the full organization"),
        ],
    },
    {
        "name": "Medical Rep 360",
        "enablements": [
            ("managers", "review rep performance — calls, coverage, LCF/RCF/MCF frequency, CLM %, coaching scores"),
            ("executives", "identify top performers and underperformers across business units"),
        ],
    },
    {
        "name": "Reports Hub",
        "enablements": [
            ("managers", "navigate to Working Days Analysis, MR 360, and Pharmacy Sales dashboards from one tile gallery"),
        ],
    },
    {
        "name": "Pharmacy Sell-Out Analytics",
        "enablements": [
            ("managers", "view distributor withdrawal data, ROI trends, and governorate-level coverage heat maps"),
            ("executives", "access AI-driven planning recommendations (Agentforce) linking sell-out to field actions"),
            ("product managers", "import IbnSina and Pharmaoverseas CSV feeds with pharmacy matching and validation"),
        ],
    },
    {
        "name": "Admin Console",
        "enablements": [
            ("managers", "configure CLM, territories, bricks, products, plan cycles, coaching templates, and sales data (SFE/RM)"),
            ("product managers", "manage product catalog, PTA alignment, rating layouts, and CLM content lifecycle"),
        ],
    },
    {
        "name": "Territory Alignment & Bricks",
        "enablements": [
            ("managers", "manage ATF/ATPF/PTA targeting, IQVIA IMS bricks, and pharmacy brick membership"),
            ("product managers", "align products to territories and control which products reps can detail"),
        ],
    },
    {
        "name": "Executive Home",
        "enablements": [
            ("executives", "view org KPI snapshots, visit trends, BU cards, workforce roster, and top performers"),
        ],
    },
    {
        "name": "Project Management",
        "enablements": [
            ("executives", "track product launches and campaigns — budget vs spend, milestones, KPIs, account goals"),
            ("product managers", "create projects, assign teams, and link field activities to launch execution"),
        ],
    },
    {
        "name": "Promo Budget",
        "enablements": [
            ("executives", "monitor promotional spend allocation and utilization by business unit"),
            ("product managers", "link promo budget lines to launch projects and field visit execution"),
        ],
    },
    {
        "name": "Cross-Department Collaboration",
        "enablements": [
            ("executives", "coordinate requests across Promo, Sales, and Medical Affairs department islands"),
            ("MSLs", "collaborate cross-functionally via the Medical Affairs department island"),
        ],
    },
    {
        "name": "Product Catalog",
        "enablements": [
            ("product managers", "manage the pharma product portfolio and brand hierarchy"),
            ("reps", "view territory-aligned products available for detailing during visits"),
        ],
    },
    {
        "name": "Employee Plan Cycle",
        "enablements": [
            ("reps", "work against monthly visit targets set in employee time cards"),
            ("managers", "manage plan cycles, coverage targets, and copy plans between months via Plan Manager"),
        ],
    },
    {
        "name": "Agentforce & Einstein Insights",
        "enablements": [
            ("reps", "receive Next Best Customer recommendations and visit context briefs"),
            ("managers", "apply AI-driven pharmacy sell-out recommendations to field planning"),
            ("executives", "access planning vision sessions and sell-out-to-action orchestration"),
        ],
    },
    {
        "name": "Integrations",
        "enablements": [
            ("managers", "monitor connectors for maps, IMS bricks, IbnSina, Pharmaoverseas, and Mendix via Admin Console"),
            ("executives", "view Mendix sync logs and external platform integration status"),
        ],
    },
]

PERSONA_LABELS = {
    "reps": "Enable reps to",
    "managers": "Enable managers to",
    "executives": "Enable executives to",
    "MSLs": "Enable MSLs to",
    "product managers": "Enable product managers to",
}

TIMELINE_ROWS = [
    ("1. Discovery", "4 weeks", "Zeta-specific requirements workshops; territory/product/CLM inventory; integration mapping; data migration assessment", "Signed requirements doc, gap analysis, project plan"),
    ("2. Implementation", "14 weeks", "Org configuration; module deployment (51 packages); Zeta branding; territory/ATF seed; CLM content load; Einstein/Agentforce setup; integration build", "Configured sandbox + partial production deploy"),
    ("3. UAT", "4 weeks", "Rep/manager/executive/MSL test scripts; defect triage; compliance sign-off", "UAT sign-off, defect closure report"),
    ("4. Training", "3 weeks", "Role-based training (5 personas); admin/SFE train-the-trainer; quick-reference guides", "Training completion records, admin runbook"),
    ("5. Go Live", "1 week", "Production cutover; hypercare war room; go/no-go checklist", "Live production org, hypercare log"),
    ("6. Support", "1 / 3 / 5 years", "Quarterly health checks; bug fixes; minor enhancements; Salesforce release prep", "SLA-backed support (see Investment Summary)"),
]

SUPPORT_ROWS = [
    ("1 year", "EGP 1,000,000 / year", "—"),
    ("3 years", "EGP 750,000 / year", "25% discount"),
    ("5 years", "EGP 600,000 / year", "40% discount"),
]

PAYMENT_SCHEDULE_ROWS = [
    ("1st", "Project kickoff", "25%"),
    ("2nd", "Discovery sign-off", "25%"),
    ("3rd", "UAT sign-off", "25%"),
    ("4th", "Go-live", "25%"),
]

IMPLEMENTATION_COST = "EGP 6,240,000"

# Split modules across two scope pages (~half each)
SCOPE_PAGE1_COUNT = 15


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(9)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)


def set_paragraph_spacing(paragraph, before=0, after=2, line=10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)


def add_cover_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Technical Proposal")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    set_paragraph_spacing(title, after=6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Zeta Pharma Commercial Platform")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = BLUE
    set_paragraph_spacing(subtitle, after=12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Prepared for: Zeta Pharma\n"
        f"Submitted by: Cloudastick Systems — Salesforce Partner\n"
        f"{date.today():%B %Y}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = GRAY
    set_paragraph_spacing(meta, after=8)


def add_scope_header(doc: Document) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run("Technical Scope — Modules & Enablement Points")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    set_paragraph_spacing(heading, before=0, after=4)


def add_module_block(doc: Document, module: dict) -> None:
    title = doc.add_paragraph()
    run = title.add_run(module["name"])
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    set_paragraph_spacing(title, before=3, after=1)

    for persona_key, text in module["enablements"]:
        p = doc.add_paragraph(style="List Bullet")
        label = PERSONA_LABELS[persona_key]
        label_run = p.add_run(f"{label} {text}")
        label_run.font.size = Pt(7.5)
        set_paragraph_spacing(p, after=0, line=9)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: List[str],
    rows: List[Tuple],
    col_widths: Optional[List[float]] = None,
    header_fill: str = "1A365D",
    header_font_pt: float = 8,
    body_font_pt: float = 8,
    bold_rows: Optional[List[int]] = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    bold_rows = bold_rows or []

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_fill)
        for paragraph in hdr_cells[i].paragraphs:
            set_paragraph_spacing(paragraph, after=0, line=11)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(header_font_pt)
                if header_fill == "1A365D":
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        is_bold = row_idx in bold_rows
        for col_idx, value in enumerate(row_data):
            row_cells[col_idx].text = value
            for paragraph in row_cells[col_idx].paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=11)
                for run in paragraph.runs:
                    run.font.size = Pt(body_font_pt)
                    if is_bold:
                        run.bold = True

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)


def add_investment_summary_page(doc: Document) -> None:
    doc.add_page_break()

    heading = doc.add_paragraph()
    run = heading.add_run("Investment Summary")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    set_paragraph_spacing(heading, after=10)

    cost_rows = [
        (
            "Implementation:\nDiscovery, Implementation, UAT, Training, Go-Live",
            IMPLEMENTATION_COST,
        ),
        (
            "Annual Support (1 Year)\nLonger commitment = lower annual rate",
            "EGP 1,000,000 / year",
        ),
        ("Annual Support (3 Years)", "EGP 750,000 / year"),
        ("Annual Support (5 Years)", "EGP 600,000 / year"),
        ("Total Implementation", IMPLEMENTATION_COST),
    ]
    add_table(
        doc,
        ["Item", "Estimated Cost"],
        cost_rows,
        col_widths=[4.25, 1.75],
        header_fill="D9D9D9",
        header_font_pt=9,
        body_font_pt=9,
        bold_rows=[5],
    )

    doc.add_paragraph()
    schedule_heading = doc.add_paragraph()
    schedule_run = schedule_heading.add_run("Payment Schedule")
    schedule_run.bold = True
    schedule_run.font.size = Pt(11)
    schedule_run.font.color.rgb = NAVY
    set_paragraph_spacing(schedule_heading, before=8, after=6)

    add_table(
        doc,
        ["Installment", "Milestone", "Percentage"],
        PAYMENT_SCHEDULE_ROWS,
        col_widths=[1.25, 3.25, 1.5],
        header_fill="D9D9D9",
        header_font_pt=9,
        body_font_pt=9,
    )

    semi_annual = doc.add_paragraph(
        "Alternative: Semi-annual payments — 50% on project kickoff, 50% at go-live."
    )
    set_paragraph_spacing(semi_annual, before=6, after=4, line=10)
    for run in semi_annual.runs:
        run.font.size = Pt(8.5)
        run.italic = True

    support_note = doc.add_paragraph(
        "Support fees are invoiced annually in advance, beginning at go-live. "
        "Support includes bug fixes, Salesforce release regression, up to 40 support hours/year, "
        "quarterly health reviews, and P1 response within 4 business hours."
    )
    set_paragraph_spacing(support_note, before=2, after=4, line=10)
    for run in support_note.runs:
        run.font.size = Pt(8)

    excludes = doc.add_paragraph(
        "Excludes: Salesforce platform licenses, third-party data feeds (IQVIA/OneKey), "
        "and travel outside Cairo."
    )
    set_paragraph_spacing(excludes, after=10, line=10)
    for run in excludes.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

    signature = doc.add_paragraph()
    sig_run = signature.add_run(
        "Accepted by: ___________________     Date: ___________\n"
        "Cloudastick Systems — Salesforce Partner"
    )
    sig_run.font.size = Pt(9)
    set_paragraph_spacing(signature, before=4)


def add_timeline_page(doc: Document) -> None:
    doc.add_page_break()

    heading = doc.add_paragraph()
    run = heading.add_run("Delivery Timeline & Phases")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    set_paragraph_spacing(heading, after=6)

    intro = doc.add_paragraph(
        "Total elapsed time from discovery kickoff to go-live: approximately 26 weeks (~6 months). "
        "Support begins at go-live and continues for the selected term (1, 3, or 5 years)."
    )
    set_paragraph_spacing(intro, after=8, line=11)
    for run in intro.runs:
        run.font.size = Pt(9)

    add_table(
        doc,
        ["Phase", "Duration", "Key Activities", "Deliverables"],
        TIMELINE_ROWS,
        col_widths=[1.1, 0.75, 2.5, 1.65],
    )

    doc.add_paragraph()
    phases = doc.add_paragraph()
    phases.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bar = (
        "Discovery (4w)  →  Implementation (14w)  →  UAT (4w)  →  "
        "Training (3w)  →  Go Live (1w)  →  Support (1/3/5 yr)"
    )
    bar_run = phases.add_run(bar)
    bar_run.bold = True
    bar_run.font.size = Pt(8)
    bar_run.font.color.rgb = BLUE
    set_paragraph_spacing(phases, before=6, after=0)


def add_commercials_page(doc: Document) -> None:
    add_investment_summary_page(doc)


def build_document() -> Document:
    doc = Document()
    set_document_defaults(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    add_cover_block(doc)
    add_scope_header(doc)

    for module in MODULES[:SCOPE_PAGE1_COUNT]:
        add_module_block(doc, module)

    doc.add_page_break()
    add_scope_header(doc)

    for module in MODULES[SCOPE_PAGE1_COUNT:]:
        add_module_block(doc, module)

    add_timeline_page(doc)
    add_commercials_page(doc)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
