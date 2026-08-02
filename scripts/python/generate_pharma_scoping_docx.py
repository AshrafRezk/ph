#!/usr/bin/env python3
"""Generate Google Docs–friendly Pharma Discovery & Scoping DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "Plan" / "Pharma_Discovery_Scoping_Template.docx"

CB = "☐"  # Unicode checkbox — converts cleanly in Google Docs


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x5C, 0xAB)


def p(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def blank_lines(doc: Document, n: int = 3) -> None:
    for _ in range(n):
        doc.add_paragraph("_" * 95)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "0B5CAB") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr[i].text = title
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(hdr[i], header_fill)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for paragraph in cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def fill_row(label: str, unit: str = "", width_hint: str = "______________") -> list[str]:
    return [label, width_hint, unit]


def module_row(name: str, description: str) -> list[str]:
    return [
        f"{CB} Include",
        name,
        description,
        f"{CB} M  {CB} S  {CB} C  {CB} W",
        "",
    ]


def main() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Salesforce Sales Cloud for Pharma")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x5C, 0xAB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Discovery & Scoping Workbook")
    r.bold = True
    r.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Fillable template for customer workshops · Upload to Google Docs · "
        "Tick ☐ boxes · Enter numbers · Capture comments"
    ).italic = True

    p(doc, "Document control", bold=True)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Customer / Company", "________________________________"],
            ["Workshop date", "____ / ____ / ________"],
            ["Prepared by", "________________________________"],
            ["Customer sponsor", "________________________________"],
            ["Target go-live", "____ / ____ / ________"],
            ["Version", "1.0 — July 2026"],
        ],
    )

    # ------------------------------------------------------------------ A
    h(doc, "1. Company & Commercial Footprint (numbers)", 1)
    p(
        doc,
        "Enter volumes used for licensing, effort, and data migration sizing. "
        "Leave blank if unknown — mark TBD.",
    )

    h(doc, "1.1 Organization scale", 2)
    add_table(
        doc,
        ["Metric", "Number", "Notes / unit"],
        [
            fill_row("Number of Business Units / Lines", "______________", "e.g. GIT, Diabetes, CHC"),
            fill_row("Number of countries / markets", "______________", ""),
            fill_row("Number of territories", "______________", ""),
            fill_row("Number of bricks (if used)", "______________", ""),
            fill_row("Number of products / SKUs promoted", "______________", ""),
            fill_row("Number of brands", "______________", ""),
        ],
    )

    h(doc, "1.2 Field force headcount", 2)
    add_table(
        doc,
        ["Role", "Number", "Notes"],
        [
            fill_row("Medical / Sales Representatives", "______________", ""),
            fill_row("MSLs / Medical Affairs field", "______________", ""),
            fill_row("District / First-line Managers (DMs)", "______________", ""),
            fill_row("Regional / Area Managers", "______________", ""),
            fill_row("SFE / Sales Ops / Commercial Ops", "______________", ""),
            fill_row("Marketing / Brand users", "______________", ""),
            fill_row("C-level / Executive users", "______________", ""),
            fill_row("Total Salesforce users (estimate)", "______________", ""),
        ],
    )

    h(doc, "1.3 Customer master & visit universe", 2)
    add_table(
        doc,
        ["Metric", "Number", "Notes"],
        [
            fill_row("Master database accounts (total)", "______________", "HCP + Pharmacy + HCO + other"),
            fill_row("HCPs in master", "______________", ""),
            fill_row("Pharmacies in master", "______________", ""),
            fill_row("Institutions / HCOs in master", "______________", ""),
            fill_row("Accounts actively visited (in-plan)", "______________", "per cycle"),
            fill_row("Avg visits per rep per day", "______________", ""),
            fill_row("Avg visits per rep per month", "______________", ""),
            fill_row("Plan cycle length", "______________", "e.g. calendar month"),
        ],
    )

    h(doc, "1.4 Revenue & commercial value (optional but useful)", 2)
    add_table(
        doc,
        ["Metric", "Amount / Number", "Currency / Notes"],
        [
            fill_row("Annual revenue in scope (total)", "______________", "________"),
            fill_row("Revenue per line / BU (list below)", "—", "see table"),
            fill_row("Avg revenue per territory (if known)", "______________", "________"),
            fill_row("Sample spend / year (approx.)", "______________", "________"),
            fill_row("Promo / marketing budget in scope", "______________", "________"),
        ],
    )

    p(doc, "Revenue by line / Business Unit (add rows as needed)", bold=True)
    add_table(
        doc,
        ["Line / BU name", "Annual revenue", "Currency", "# Reps on line", "# Products", "Priority (H/M/L)"],
        [
            ["________________", "______________", "______", "________", "________", "☐ H  ☐ M  ☐ L"],
            ["________________", "______________", "______", "________", "________", "☐ H  ☐ M  ☐ L"],
            ["________________", "______________", "______", "________", "________", "☐ H  ☐ M  ☐ L"],
            ["________________", "______________", "______", "________", "________", "☐ H  ☐ M  ☐ L"],
            ["________________", "______________", "______", "________", "________", "☐ H  ☐ M  ☐ L"],
        ],
    )

    h(doc, "1.5 Current systems", 2)
    add_table(
        doc,
        ["System area", "Current tool", "Keep / Replace / Integrate"],
        [
            ["CRM / Call reporting", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
            ["CLM / e-detailing", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
            ["Sample management", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
            ["Master data (OneKey / IQVIA / other)", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
            ["Wholesaler / sell-out data", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
            ["HR / TOT / leave", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
            ["ERP / finance", "________________________", f"{CB} Keep  {CB} Replace  {CB} Integrate"],
        ],
    )

    p(doc, "Comments — footprint, licensing assumptions, exclusions", bold=True)
    blank_lines(doc, 4)

    # ------------------------------------------------------------------ B
    h(doc, "2. Module Scope — Include & Priority", 1)
    p(
        doc,
        "For each module: tick Include, then set MoSCoW priority "
        "(M = Must go-live · S = Should wave 1–2 · C = Could later · W = Won’t this program). "
        "Use Comments for customization or business rules.",
    )

    add_table(
        doc,
        ["Include", "Module", "What it covers", "Priority", "Comments / customization"],
        [
            module_row(
                "1. Account Foundation",
                "HCP / Pharmacy / HCO / BC record types, specialties, layouts",
            ),
            module_row(
                "2. Territory & Targeting",
                "ATF / ATPF / PTA classification, product alignment",
            ),
            module_row(
                "3. Bricks & Geography",
                "Brick master, pharmacy–brick links, planning cells",
            ),
            module_row(
                "4. Field Rep Home",
                "Coverage / RF% KPIs, today route map, Next Best Customer",
            ),
            module_row(
                "5. Accounts Tab",
                "Territory account list, filters, map pins, collections",
            ),
            module_row(
                "6. Field Rep Planner",
                "Week calendar, drag-drop visits, TOT, map + route optimize",
            ),
            module_row(
                "7. Visit / Call Reporting",
                "Attendees, products, messages, samples, CLM, mobile + desktop",
            ),
            module_row(
                "8. Sample Management",
                "Lot / expiry inventory, issue on visit, audit transactions",
            ),
            module_row(
                "9. CLM (Closed Loop Marketing)",
                "Content library, player, dwell time, message feedback, ratings",
            ),
            module_row(
                "10. Product Surveys / Messaging",
                "HCP survey links / WhatsApp-style engagement from visit",
            ),
            module_row(
                "11. Medical Inquiry",
                "Capture MI from visit → Case / Medical Affairs queue",
            ),
            module_row(
                "12. Account Affiliations",
                "HCP↔Pharmacy↔HCO network graph & attendee discovery",
            ),
            module_row(
                "13. Account Ratings",
                "Configurable rating forms; optional AI validity check",
            ),
            module_row(
                "14. Coaching & Development",
                "Templates, dual scoring, ride-along / double-visit coaching",
            ),
            module_row(
                "15. Time Off Territory (TOT)",
                "Submit / approve leave & non-field days; coverage honesty",
            ),
            module_row(
                "16. Plan Cycle & Medical Rep 360",
                "Monthly targets, coverage / RF / CLM% / coaching metrics",
            ),
            module_row(
                "17. Pharmacy Sell-Out Analytics",
                "Wholesaler CSV import, brick dashboards, planning insights",
            ),
            module_row(
                "18. Agentforce / AI Insights",
                "NBC, pharmacy recommendations, account / visit narratives",
            ),
            module_row(
                "19. Executive Projects",
                "Campaigns, budgets, milestones, KPIs, visit linkage",
            ),
            module_row(
                "20. Promo Budget & Collaboration",
                "Promo utilization, cross-department collaboration requests",
            ),
            module_row(
                "21. Management / Executive Dashboards",
                "Team KPI command center, reports hub, heat maps",
            ),
            module_row(
                "22. Field Force Tracking (Fleet)",
                "Rep location publish; manager live fleet map",
            ),
            module_row(
                "23. Home Office Messages",
                "HQ announcements on Field Home",
            ),
            module_row(
                "24. Admin Console",
                "Self-service CLM, territory, bricks, plan, coaching ops",
            ),
            module_row(
                "25. Product Catalog",
                "Therapy / brand catalog, images, territory product feeds",
            ),
            module_row(
                "26. Mobile Offline (Hybrid)",
                "CLM + visits + planner offline (Briefcase + device cache)",
            ),
            module_row(
                "27. Integrations Scaffold",
                "Maps/OSRM, CSV, MDM/ERP/email connectors as required",
            ),
        ],
        header_fill="0B5CAB",
    )

    p(doc, "Module scope — free-form business explanation / exclusions", bold=True)
    blank_lines(doc, 5)

    # ------------------------------------------------------------------ C offline
    h(doc, "3. Offline Journeys (must work without network?)", 1)
    p(doc, "Tick each journey that must work offline on Salesforce Mobile. Add device / sync notes.")
    add_table(
        doc,
        ["Required offline?", "Journey", "Priority (1–5)", "Comments"],
        [
            [f"{CB} Yes   {CB} No", "Play CLM + capture session / feedback", "____", ""],
            [f"{CB} Yes   {CB} No", "Open / complete visit call report", "____", ""],
            [f"{CB} Yes   {CB} No", "View & edit weekly planner", "____", ""],
            [f"{CB} Yes   {CB} No", "Coaching create / submit from visit", "____", ""],
            [f"{CB} Yes   {CB} No", "Field Home today plan / NBC (read or draft)", "____", ""],
            [f"{CB} Yes   {CB} No", "Sample lines on visit (offline caveats OK?)", "____", ""],
        ],
    )
    p(doc, "Offline licensing / conflict / prefetch comments", bold=True)
    blank_lines(doc, 3)

    # ------------------------------------------------------------------ D discovery
    h(doc, "4. Discovery Checklist — Business Rules", 1)
    p(doc, "Tick topics discussed. Capture decisions in Comments.")

    h(doc, "4.1 Customer master", 2)
    add_table(
        doc,
        ["Done?", "Question / topic", "Decision / comments"],
        [
            [CB, "Account types Day 1 (HCP / Pharmacy / HCO / BC / other)", ""],
            [CB, "Person Accounts for HCPs? Multi-specialty fields?", ""],
            [CB, "Master data source & sync cadence", ""],
            [CB, "Affiliations required for attendees / sell-out?", ""],
            [CB, "Rating cadence (potential / loyalty refresh)", ""],
        ],
    )

    h(doc, "4.2 Targeting & visits", 2)
    add_table(
        doc,
        ["Done?", "Question / topic", "Decision / comments"],
        [
            [CB, "Classification matrix & visit frequency rules", ""],
            [CB, "Who maintains ATF / ATPF / PTA each cycle?", ""],
            [CB, "Required visit sections & Completed = read-only?", ""],
            [CB, "Manager approval on visit complete?", ""],
            [CB, "Desktop + mobile both required?", ""],
        ],
    )

    h(doc, "4.3 Compliance (CLM / samples / MI)", 2)
    add_table(
        doc,
        ["Done?", "Question / topic", "Decision / comments"],
        [
            [CB, "CLM formats & MLR / medical approval gate", ""],
            [CB, "Dwell / message / rating analytics required?", ""],
            [CB, "Sample lot / expiry / signature / transfer rules", ""],
            [CB, "Medical Inquiry queue / SLA / owners", ""],
            [CB, "HCP survey / messaging consent requirements", ""],
        ],
    )

    h(doc, "4.4 People & leadership", 2)
    add_table(
        doc,
        ["Done?", "Question / topic", "Decision / comments"],
        [
            [CB, "Coaching competency model & dual scoring", ""],
            [CB, "TOT types & approval path; impact on coverage", ""],
            [CB, "Plan-cycle KPIs for Medical Rep 360", ""],
            [CB, "Fleet GPS tracking — privacy / consent", ""],
            [CB, "Executive projects / promo budgets in scope?", ""],
        ],
    )

    p(doc, "Additional business rules / customization requests", bold=True)
    blank_lines(doc, 5)

    # ------------------------------------------------------------------ E waves
    h(doc, "5. Delivery Waves (proposed)", 1)
    p(doc, "Tick agreed waves and list modules from Section 2.")
    add_table(
        doc,
        ["Include", "Wave", "Focus", "Modules (numbers / names)", "Target date"],
        [
            [CB, "Wave 0", "Foundation — accounts, products, territory, security, Admin Console", "", "____/____/____"],
            [CB, "Wave 1", "Field core — Home, Planner, Visits, Samples, Offline", "", "____/____/____"],
            [CB, "Wave 2", "Content & medical — CLM, MI, surveys", "", "____/____/____"],
            [CB, "Wave 3", "People & cycle — Coaching, TOT, Plan 360", "", "____/____/____"],
            [CB, "Wave 4", "Market & AI — sell-out, Agentforce", "", "____/____/____"],
            [CB, "Wave 5", "Leadership — projects, budgets, dashboards", "", "____/____/____"],
        ],
    )

    p(doc, "Phasing comments / dependencies / pilot territories", bold=True)
    blank_lines(doc, 4)

    # ------------------------------------------------------------------ F sizing summary
    h(doc, "6. Quick Sizing Summary (copy to proposal)", 1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["# Sales / Medical Reps", "______________"],
            ["# District Managers", "______________"],
            ["# Lines / BUs", "______________"],
            ["Revenue in scope (total)", "______________  ________"],
            ["Master accounts (total)", "______________"],
            ["Accounts visited (in-plan)", "______________"],
            ["Modules Included (count)", "______________"],
            ["Must-have modules (list)", "______________________________________________"],
            ["Pilot user count", "______________"],
            ["Pilot go-live", "____ / ____ / ________"],
        ],
    )

    # ------------------------------------------------------------------ G sign-off
    h(doc, "7. Sign-off", 1)
    add_table(
        doc,
        ["Role", "Name", "Date", "Signature / Initials"],
        [
            ["Customer sponsor", "", "", ""],
            ["Commercial / SFE", "", "", ""],
            ["IT / Salesforce owner", "", "", ""],
            ["Compliance / Medical (if needed)", "", "", ""],
            ["Delivery / SI lead", "", "", ""],
        ],
    )

    h(doc, "8. Appendix — Salesforce value (talk track)", 1)
    p(
        doc,
        "One platform for the full commercial loop: targeted accounts → plan → visit → "
        "CLM & samples → coaching & TOT → sell-out insights → leadership dashboards — "
        "with mobile offline for CLM, visit logging, and planner. Configurable via Admin Console "
        "each cycle without waiting on IT for every change.",
        italic=True,
    )

    p(doc, "Open risks / parking lot", bold=True)
    blank_lines(doc, 4)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
