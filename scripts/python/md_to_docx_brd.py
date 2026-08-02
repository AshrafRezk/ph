# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""Convert Plan/Zeta_Pharma_BRD.md to a Google Docs-friendly DOCX."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "Plan" / "Zeta_Pharma_BRD.md"
OUT = ROOT / "Plan" / "Zeta_Pharma_BRD.docx"

HEADING_BLUE = RGBColor(0x0B, 0x5C, 0xAB)
HEADER_FILL = "0B5CAB"
QUOTE_FILL = "F3F6F9"

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def add_runs(paragraph, text: str, *, base_bold: bool = False, base_italic: bool = False) -> None:
    if not text:
        return
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = base_bold
            run.italic = base_italic
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.bold = base_bold
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.bold = base_bold
            run.italic = base_italic
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            label, _url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(label)
            run.bold = base_bold
            run.italic = base_italic
            run.font.color.rgb = HEADING_BLUE
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic


def add_paragraph(doc: Document, text: str, *, italic: bool = False, quote: bool = False) -> None:
    para = doc.add_paragraph()
    if quote:
        para.paragraph_format.left_indent = Inches(0.2)
    add_runs(para, text, base_italic=italic or quote)
    if quote:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_heading(doc: Document, text: str, level: int) -> None:
    # Strip markdown emphasis from headings for clean TOC text
    clean = re.sub(r"[*_`]", "", text).strip()
    p = doc.add_heading(clean, level=min(level, 4))
    for run in p.runs:
        run.font.color.rgb = HEADING_BLUE


def add_list_item(doc: Document, text: str, ordered: bool = False, number: int | None = None) -> None:
    style = "List Number" if ordered else "List Bullet"
    para = doc.add_paragraph(style=style)
    # python-docx list styles already number/bullet; keep text only
    add_runs(para, text)


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        # skip separator row
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            i += 1
            continue
        rows.append(cells)
        i += 1
    headers = rows[0]
    body = rows[1:]
    # normalize column counts
    width = len(headers)
    body = [r + [""] * (width - len(r)) if len(r) < width else r[:width] for r in body]
    return headers, body, i


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, title, base_bold=True)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        set_cell_shading(cell, HEADER_FILL)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, val)
            for run in p.runs:
                run.font.size = Pt(9)
            # light zebra
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F7FAFC")
    doc.add_paragraph()


def convert(md_text: str) -> Document:
    doc = Document()
    style_doc(doc)

    lines = md_text.splitlines()
    i = 0
    title_done = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # ATX headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            if level == 1 and not title_done:
                p = doc.add_heading(re.sub(r"[*_`]", "", heading_text), level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.color.rgb = HEADING_BLUE
                title_done = True
            else:
                add_heading(doc, heading_text, level)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-{3,}", lines[i + 1].strip()):
            headers, body, i = parse_table(lines, i)
            add_table(doc, headers, body)
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_parts: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_parts.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            add_paragraph(doc, " ".join(quote_parts).strip() or " ", quote=True)
            continue

        # Unordered list
        if re.match(r"^[-*+]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*+]\s+", lines[i].strip()):
                item = re.sub(r"^[-*+]\s+", "", lines[i].strip())
                add_list_item(doc, item, ordered=False)
                i += 1
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                add_list_item(doc, item, ordered=True)
                i += 1
            continue

        # Metadata-ish bold lines under title (**Project:** ...)
        if stripped.startswith("**") and ":**" in stripped:
            add_paragraph(doc, stripped)
            i += 1
            continue

        # Normal paragraph (merge continuation lines that aren't special)
        para_parts = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith(">") or re.match(r"^[-*+]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            para_parts.append(nxt)
            i += 1
        add_paragraph(doc, " ".join(para_parts))

    return doc


def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = convert(md)
    # Cover note at end of first section isn't needed; save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
